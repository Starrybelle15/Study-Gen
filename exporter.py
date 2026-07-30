"""
QGen AI
exporter.py
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt


def export_docx(questions):

    doc = Document()

    title = doc.add_heading("StudyGen AI", level=1)
    title.style.font.size = Pt(22)

    doc.add_paragraph(
        "Automatically Generated Study Questions"
    )

    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}"
    )

    doc.add_paragraph("-" * 60)

    for line in questions.split("\n"):

        if line.strip():

            doc.add_paragraph(line)

    filename = "QGen_Questions.docx"

    doc.save(filename)

    return filename